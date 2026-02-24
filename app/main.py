from docx import Document
from fastapi import FastAPI, UploadFile, File, Request, HTTPException
from fastapi.templating import Jinja2Templates
from fastapi.staticfiles import StaticFiles
from fastapi.responses import JSONResponse
import csv, uuid, json, qrcode, os, base64, random, string
import io
from io import BytesIO
from dotenv import load_dotenv
import psycopg2
from psycopg2.extras import Json, RealDictCursor
from fastapi import Form
import uuid
from docxtpl import DocxTemplate
import subprocess
from reportlab.pdfgen import canvas
from reportlab.lib.pagesizes import A4
from pypdf import PdfReader, PdfWriter
from fastapi.responses import FileResponse, RedirectResponse
from fastapi.responses import StreamingResponse
# from PyPDF2 import PdfReader, PdfWriter
from reportlab.lib.pagesizes import letter
from reportlab.lib.utils import ImageReader



# ------------------ LOAD ENV ------------------
load_dotenv(dotenv_path=os.path.join(os.path.dirname(__file__), "..", ".env"))

DATABASE_URL = os.getenv("DATABASE_URL")
BASE_URL = os.getenv("BASE_URL", "http://localhost:8000")

if not DATABASE_URL:
    raise Exception("DATABASE_URL is missing in .env")

print("DATABASE_URL loaded successfully!")


# ------------------ APP SETUP ------------------
app = FastAPI()
app.mount("/static", StaticFiles(directory="app/static"), name="static")
templates = Jinja2Templates(directory="app/templates")
app.mount("/certificates", StaticFiles(directory="certificates"), name="certificates")


CERT_DIR = "certificates"
os.makedirs(CERT_DIR, exist_ok=True)
os.makedirs("certificates", exist_ok=True)

BASE_DIR = os.path.dirname(os.path.abspath(__file__))

TEMPLATE_CERT_PATH = os.path.join(BASE_DIR, "templates", "ONBOARDING_FINAL.pdf")


# ------------------ HELPERS ------------------
def generate_url_segments():
    param1 = str(random.randint(100, 999))  # e.g., 566
    param2 = "".join(random.choices(string.ascii_lowercase, k=7))  # e.g., aprivate
    param3 = str(random.randint(1, 10))  # e.g., 4
    param4 = str(random.randint(100, 999))  # e.g., 243
    param5 = str(random.randint(10, 99))  # e.g., 32 (extra for your example)
    return param1, param2, param3, param4, param5


def normalize_remark(value: str) -> str:
    value = value.strip().lower()
    value = value.replace("recomended", "recommended")
    return value


def count_recommended(rows: list) -> int:
    return sum(
        1
        for r in rows
        if "recommended" in r.get("remark", "") and "not" not in r.get("remark", "")
    )


def get_verification_context(request: Request, record_id: str):
    with get_db() as conn:
        with conn.cursor() as cur:
            cur.execute("SELECT * FROM uploads WHERE record_id = %s", (record_id,))
            row = cur.fetchone()

    if not row:
        raise HTTPException(status_code=404, detail="Record not found")

    rows = row["rows"]
    recommended_count = count_recommended(rows)

    qr_data = (
        f"{BASE_URL}/wassce-list/"
        f"{row['param1']}/{row['param2']}/{row['param3']}/"
        f"{row['param4']}/{row['param5']}/{record_id}"
    )

    qr_img = qrcode.make(qr_data)
    buf = BytesIO()
    qr_img.save(buf, format="PNG")
    qr_base64 = base64.b64encode(buf.getvalue()).decode()

    return {
        "request": request,
        "school_name": row["school_name"],
        "school_code": row["school_code"],
        "principal": row["principal"],
        "recommended_count": recommended_count,
        "qr_code": qr_base64,
        "rows": rows,  # ✅ THIS WAS MISSING
    }


# ------------------ DATABASE ------------------
def get_db():
    return psycopg2.connect(DATABASE_URL, cursor_factory=RealDictCursor)


def init_db():
    with get_db() as conn:
        with conn.cursor() as cur:
            cur.execute("""
                CREATE TABLE IF NOT EXISTS uploads (
                    record_id TEXT PRIMARY KEY,
                    school_name TEXT,
                    school_code TEXT,
                    principal TEXT,
                    rows JSONB,
                    param1 TEXT,
                    param2 TEXT,
                    param3 TEXT,
                    param4 TEXT,
                    param5 TEXT
                )
            """)
            cur.execute("""
            CREATE TABLE IF NOT EXISTS onboarding_certificates (
                id SERIAL PRIMARY KEY,
                record_id TEXT UNIQUE,
                school_name TEXT,
                lga TEXT,
                school_code TEXT,
                js1 INTEGER,
                js2 INTEGER,
                js3 INTEGER,
                ss1 INTEGER,
                ss2 INTEGER,
                ss3 INTEGER,
                total INTEGER,
                teachers INTEGER,
                pdf_path TEXT,
                Our_Ref INTEGER,
                date_of_issue DATE,
                created_at TIMESTAMP DEFAULT CURRENT_TIMESTAMP
            )
            """)


init_db()


# ------------------ ROUTES ------------------
@app.get("/")
def upload_page(request: Request):
    with get_db() as conn:
        with conn.cursor() as cur:
            cur.execute(
                "SELECT * FROM onboarding_certificates ORDER BY created_at DESC"
            )
            records = cur.fetchall()

    return templates.TemplateResponse(
        "upload.html",
        {"request": request, "records": records},
    )


@app.post("/upload")
async def upload_csv(file: UploadFile = File(...)):
    content = (await file.read()).decode("utf-8").splitlines()

    data_rows = []
    school_name = ""
    school_code = ""
    principal = ""
    header_found = False

    # WAEC-style CSV
    if any("SCHOOL NAME" in line.upper() for line in content[:10]):
        reader = csv.reader(content)
        for row in reader:
            if not row:
                continue
            if row[0].strip().upper() == "SCHOOL NAME":
                school_name = row[1].strip()
                continue
            if row[0].strip().upper() == "SCHOOL CODE":
                school_code = row[1].strip()
                continue
            if row[0].strip().upper() == "NAME OF PRINCIPAL":
                principal = row[1].strip()
                continue
            if row[0].strip() == "#":
                header_found = True
                continue
            if not header_found:
                continue
            data_rows.append(
                {
                    "passport": row[1],
                    "lin": row[2],
                    "lastname": row[3],
                    "firstname": row[4],
                    "othername": row[5],
                    "sex": row[6],
                    "year_2026": row[7],
                    "year_2025": row[8],
                    "year_2024": row[9],
                    "remark": normalize_remark(row[10]),
                }
            )
    else:
        reader = csv.DictReader(content)
        for row in reader:
            row["remark"] = normalize_remark(row.get("remark", ""))
            data_rows.append(row)
        school_name = "CLIMAX SECONDARY SCHOOL"
        school_code = "C24084"
        principal = "Tolani Ogunbamiji"

    record_id = f"{school_code}-{uuid.uuid4().hex[:6]}"

    # Generate fixed URL segments
    param1, param2, param3, param4, param5 = generate_url_segments()

    # Insert into Postgres with URL segments
    with get_db() as conn:
        with conn.cursor() as cur:
            cur.execute(
                """
                INSERT INTO uploads
                (record_id, school_name, school_code, principal, rows, param1, param2, param3, param4, param5)
                VALUES (%s, %s, %s, %s, %s, %s, %s, %s, %s, %s)
            """,
                (
                    record_id,
                    school_name,
                    school_code,
                    principal,
                    Json(data_rows),
                    param1,
                    param2,
                    param3,
                    param4,
                    param5,
                ),
            )

    custom_index_url = f"{BASE_URL}/wassce-list/{param1}/{param2}/{param3}/{param4}/{param5}/{record_id}"

    return JSONResponse(
        {
            "message": "Upload successful",
            "index_url": custom_index_url,  # ✅ customised URL
            "verify_url": f"/verify/{record_id}",  # ✅ NOT customised
        }
    )


@app.get("/verify/{record_id}")
def verify_page(request: Request, record_id: str):
    context = get_verification_context(request, record_id)
    return templates.TemplateResponse("verification.html", context)


@app.get("/wassce-list/{param1}/{param2}/{param3}/{param4}/{param5}/{record_id}")
def wassce_list_page(
    request: Request,
    param1: str,
    param2: str,
    param3: str,
    param4: str,
    param5: str,
    record_id: str,
):
    with get_db() as conn:
        with conn.cursor() as cur:
            cur.execute(
                """
                        SELECT
                        record_id,
                        param1,
                        param2,
                        param3,
                        param4,
                        param5
                        FROM uploads
                        WHERE record_id = %s
            """,
                (record_id,),
            )
            row = cur.fetchone()

    if not row:
        raise HTTPException(status_code=404, detail="Record not found")

    # Ensure the URL matches the stored segments
    if (row["param1"], row["param2"], row["param3"], row["param4"], row["param5"]) != (
        param1,
        param2,
        param3,
        param4,
        param5,
    ):
        raise HTTPException(status_code=404, detail="URL segments do not match record")

    context = get_verification_context(request, record_id)
    context.update(
        {
            "param1": param1,
            "param2": param2,
            "param3": param3,
            "param4": param4,
            "param5": param5,
        }
    )
    return templates.TemplateResponse("index.html", context)


# ------------------ONBOARDING  CERTIFICATE ROUTES ------------------


@app.post("/generate-certificate")
def generate_certificate(
    school_name: str = Form(...),
    lga: str = Form(...),
    school_code: str = Form(...),
    js1: int = Form(0),
    js2: int = Form(0),
    js3: int = Form(0),
    ss1: int = Form(0),
    ss2: int = Form(0),
    ss3: int = Form(0),
    total: int = Form(0),
    teachers: int = Form(0),
):
    record_id = uuid.uuid4().hex[:8]

    with get_db() as conn:
        with conn.cursor() as cur:
            cur.execute(
                """
                INSERT INTO onboarding_certificates
                (record_id, school_name, lga, school_code,
                 js1, js2, js3, ss1, ss2, ss3, total, teachers)
                VALUES (%s,%s,%s,%s,%s,%s,%s,%s,%s,%s,%s,%s)
            """,
                (
                    record_id,
                    school_name,
                    lga,
                    school_code,
                    js1,
                    js2,
                    js3,
                    ss1,
                    ss2,
                    ss3,
                    total,
                    teachers,
                ),
            )

    return RedirectResponse("/certificates-list", status_code=303)


@app.get("/certificates-list")
def certificates_list(request: Request):
    with get_db() as conn:
        with conn.cursor() as cur:
            cur.execute(
                "SELECT * FROM onboarding_certificates ORDER BY created_at DESC"
            )
            records = cur.fetchall()

    return templates.TemplateResponse(
        "upload.html", {"request": request, "records": records}
    )


def replace_text_preserve_layout(doc, replacements):

    def replace_in_paragraph(paragraph):
        full_text = "".join(run.text for run in paragraph.runs)
        if not any(k in full_text for k in replacements):
            return

        for run in paragraph.runs:
            for key, val in replacements.items():
                if key in run.text:
                    run.text = run.text.replace(key, val)

    for p in doc.paragraphs:
        replace_in_paragraph(p)

    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for p in cell.paragraphs:
                    replace_in_paragraph(p)


@app.post("/merge/{record_id}")
def merge_certificate(record_id: str):
    with get_db() as conn:
        with conn.cursor() as cur:
            cur.execute(
                "UPDATE onboarding_certificates SET pdf_path='READY' WHERE record_id=%s",
                (record_id,),
            )
    return {"message": "Merged"}


TEMPLATE_DOCX = "app/templates/Templatesword.docx"
CERTIFICATES_DIR = "certificates"
os.makedirs(CERTIFICATES_DIR, exist_ok=True)
LIBREOFFICE_PATH = os.getenv("LIBREOFFICE_PATH", r"C:\Program Files\LibreOffice\program\soffice.exe")

def build_certificate_word(row):
    doc = Document(TEMPLATE_DOCX)

    replacements = {
        "{{Our_Ref}}": str(row.get("Our_Ref", "")),
        "{{date_of_issue}}": str(row.get("date_of_issue", "")),
        "{{school_name}}": str(row.get("school_name", "")),
        "{{lga}}": str(row.get("lga", "")),
        "{{school_code}}": str(row.get("school_code", "")),
        "{{js1}}": str(row.get("js1", 0)),
        "{{js2}}": str(row.get("js2", 0)),
        "{{js3}}": str(row.get("js3", 0)),
        "{{ss1}}": str(row.get("ss1", 0)),
        "{{ss2}}": str(row.get("ss2", 0)),
        "{{ss3}}": str(row.get("ss3", 0)),
        "{{total}}": str(row.get("total", 0)),
        "{{teachers}}": str(row.get("teachers", 0)),
    }

    replace_text_preserve_layout(doc, replacements)

    word_path = os.path.join(CERTIFICATES_DIR, f"{row['record_id']}.docx")
    doc.save(word_path)

    return word_path



@app.get("/print/{record_id}")
def print_certificate(record_id: str):
    with get_db() as conn:
        with conn.cursor() as cur:
            cur.execute(
                "SELECT * FROM onboarding_certificates WHERE record_id=%s",
                (record_id,),
            )
            row = cur.fetchone()

    if not row:
        raise HTTPException(404, "Certificate not found")

    word_path = os.path.join(CERTIFICATES_DIR, f"{record_id}.docx")

    if not os.path.exists(word_path):
        word_path = build_certificate_word(row)

    return FileResponse(
        word_path,
        media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        filename=f"Certificate_{record_id}.docx"
    )



### pdf upload and QR code embedding routes



# Folder to store PDFs
PDF_DIR = "uploaded_pdfs"
os.makedirs(PDF_DIR, exist_ok=True)

# In-memory storage of uploaded records
uploaded_records = []

@app.post("/upload-pdf")
async def upload_pdf(file: UploadFile = File(...)):
    # 1️⃣ Save original PDF
    unique_id = uuid.uuid4().hex[:8]
    original_pdf_path = os.path.join(PDF_DIR, f"{unique_id}.pdf")
    final_pdf_path = os.path.join(PDF_DIR, f"{unique_id}_qr.pdf")
    content = await file.read()
    with open(original_pdf_path, "wb") as f:
        f.write(content)

    # 2️⃣ Generate QR code pointing to final PDF URL
    pdf_url = f"{BASE_URL}/pdf/{unique_id}_qr.pdf"
    qr = qrcode.QRCode(box_size=4, border=2)
    qr.add_data(pdf_url)
    qr.make(fit=True)
    qr_img = qr.make_image(fill_color="black", back_color="white").convert("RGB")
    qr_bytes = BytesIO()
    qr_img.save(qr_bytes, format="PNG")
    qr_bytes.seek(0)
    qr_reader = ImageReader(qr_bytes)

    # 3️⃣ Overlay QR on first page
    packet = BytesIO()
    can = canvas.Canvas(packet, pagesize=letter)
    can.drawImage(qr_reader, 450, 20, width=100, height=100)
    can.save()
    packet.seek(0)

    original_pdf = PdfReader(original_pdf_path)
    overlay_pdf = PdfReader(packet)
    writer = PdfWriter()

    for i, page in enumerate(original_pdf.pages):
        if i == 0:
            page.merge_page(overlay_pdf.pages[0])
        writer.add_page(page)

    # 4️⃣ Save final PDF with QR
    with open(final_pdf_path, "wb") as f:
        writer.write(f)

    # 5️⃣ Store record
    record = {
        "id": unique_id,
        "file_name": file.filename,
        "pdf_with_qr": f"/pdf/{unique_id}_qr.pdf"
    }
    uploaded_records.append(record)

    return JSONResponse(record)

@app.get("/pdf/{file_name}")
def serve_pdf(file_name: str):
    path = os.path.join(PDF_DIR, file_name)
    if not os.path.exists(path):
        return {"error": "PDF not found"}
    return FileResponse(path, media_type="application/pdf", filename=file_name)

@app.get("/records")
def get_records():
    return {"records": uploaded_records}